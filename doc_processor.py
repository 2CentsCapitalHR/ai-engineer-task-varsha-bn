# doc_processor.py
import re
from docx import Document
from typing import List, Dict, Any
import json
import os

def read_docx_text(path: str) -> List[str]:
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return paras

DOC_TYPE_KEYWORDS = {
    "Articles of Association": ["articles of association", "aoa", "article(s) of association"],
    "Memorandum of Association": ["memorandum of association", "moa", "memorandum"],
    "Incorporation Application Form": ["incorporation application", "application for incorporation"],
    "UBO Declaration Form": ["ubo declaration", "ultimate beneficial owner"],
    "Register of Members and Directors": ["register of members", "register of directors"]
}

def match_doc_type(text_paragraphs: List[str]) -> str:
    lower_text = "\\n".join(text_paragraphs).lower()
    scores = {}
    for dtype, kws in DOC_TYPE_KEYWORDS.items():
        score = sum(1 for kw in kws if kw in lower_text)
        scores[dtype] = score
    best = max(scores, key=lambda k: scores[k])
    if scores[best] >= 1:
        return best
    return "Unknown Document"

def rule_based_checks(paragraphs: List[str]) -> List[Dict[str, Any]]:
    issues = []
    joined = "\\n".join(paragraphs).lower()
    # jurisdiction mismatch example
    if ("federal courts" in joined or "uae federal courts" in joined) and "adgm" not in joined:
        issues.append({
            "section": "Jurisdiction",
            "issue": "References federal UAE courts rather than ADGM courts.",
            "severity": "High",
            "suggestion": "Update jurisdiction clause to explicitly reference ADGM Courts, per ADGM regulations."
        })
    # missing signatory
    if not re.search(r"signature|signed by|for and on behalf", joined):
        issues.append({
            "section": "Signatory",
            "issue": "No signatory block or signature instructions found.",
            "severity": "High",
            "suggestion": "Add signatory name, title, and signature block with date."
        })
    # permissive language
    if re.search(r"\\bmay\\b", joined) and not re.search(r"\\bshall\\b", joined):
        issues.append({
            "section": "Binding Language",
            "issue": "Document uses permissive language (e.g., 'may') which may be non-binding in critical clauses.",
            "severity": "Medium",
            "suggestion": "Where binding obligations intended, prefer 'shall' or clearly drafted obligations."
        })
    return issues

def annotate_docx(in_path: str, out_path: str, findings: List[Dict[str, Any]]):
    doc = Document(in_path)
    # Append a summary section at the end containing all comments
    doc.add_page_break()
    hdr = doc.add_paragraph()
    hdr.add_run("=== REVIEW SUMMARY ===").bold = True
    for i, f in enumerate(findings, start=1):
        p = doc.add_paragraph()
        p.add_run(f"COMMENT {i}: [{f.get('severity')}] {f.get('issue')} (section: {f.get('section')})").bold = True
        if f.get('suggestion'):
            doc.add_paragraph(f"Suggestion: {f.get('suggestion')}")
    doc.save(out_path)

def analyze_and_annotate(file_path: str, output_dir: str, checklist: Dict[str, Any]) -> Dict[str, Any]:
    paragraphs = read_docx_text(file_path)
    doc_type = match_doc_type(paragraphs)
    findings = rule_based_checks(paragraphs)
    reviewed_filename = os.path.join(output_dir, os.path.basename(file_path).replace(".docx", "_reviewed.docx"))
    annotate_docx(file_path, reviewed_filename, findings)
    report = {
        "document_name": os.path.basename(file_path),
        "document_type": doc_type,
        "issues_found": findings,
        "reviewed_file": reviewed_filename
    }
    return report
